# encoding: utf-8
import base64
import copy
import io
import os
import re

import docx
import expand
import markdown
import selektion
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import RGBColor
from myLogger import logger

paramRE = re.compile(r"\${(\w*?)}")
fmtRE = re.compile(r"\.fmt\((.*?)\)")
strokeRE = r'(\~{2})(.+?)\1'
ulRE = r'(\^{2})(.+?)\1'
STX = '\u0002'  # Use STX ("Start of text") for start-of-placeholder
ETX = '\u0003'  # Use ETX ("End of text") for end-of-placeholder
stxEtxRE = re.compile(r'%s(\d+)%s' % (STX, ETX))
headerFontSizes = [0, 24, 18, 14, 12, 10, 8]  # h1-h6 headers have fontsizes 24-8
debug = False
nlctr = 0
adfc_blue = 0x004b7c  # CMYK=90 60 10 30
adfc_yellow = 0xee7c00  # CMYK=0 60 100 0


def str2hex(s: str):
    return ":".join("{:04x}".format(ord(c)) for c in s)


def delete_paragraph(paragraph):
    # https://github.com/python-openxml/python-docx/issues/33
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def delete_run(run):
    r = run._element
    r.getparent().remove(r)
    r._p = r._element = None


def add_run_copy(paragraph, run, text=None, style=None):
    newr = paragraph.add_run(text=run.text if text is None else text,
                             style=run.style if style is None else style)
    newr.bold = run.bold
    newr.italic = run.italic
    newr.underline = run.underline
    newr.font.all_caps = run.font.all_caps
    newr.font.bold = run.font.bold
    if run.font.color is not None and run.font.color.rgb is not None:
        newr.font.color.rgb = run.font.color.rgb
    if run.font.color is not None and run.font.color.theme_color is not None:
        newr.font.color.theme_color = run.font.color.theme_color
    # newr.font.color.type = run.font.color.type is readonly
    newr.font.complex_script = run.font.complex_script
    newr.font.cs_bold = run.font.cs_bold
    newr.font.cs_italic = run.font.cs_italic
    newr.font.double_strike = run.font.double_strike
    newr.font.emboss = run.font.emboss
    newr.font.hidden = run.font.hidden
    newr.font.highlight_color = run.font.highlight_color
    newr.font.imprint = run.font.imprint
    newr.font.italic = run.font.italic
    newr.font.math = run.font.math
    newr.font.name = run.font.name
    newr.font.no_proof = run.font.no_proof
    newr.font.outline = run.font.outline
    newr.font.rtl = run.font.rtl
    newr.font.shadow = run.font.shadow
    newr.font.size = run.font.size
    newr.font.small_caps = run.font.small_caps
    newr.font.snap_to_grid = run.font.snap_to_grid
    newr.font.spec_vanish = run.font.spec_vanish
    newr.font.strike = run.font.strike
    newr.font.subscript = run.font.subscript
    newr.font.superscript = run.font.superscript
    newr.font.underline = run.font.underline
    newr.font.web_hidden = run.font.web_hidden
    return newr


def copyPara(para, newp):
    newp.alignment = para.alignment
    newp.style = para.style
    newp.paragraph_format.alignment = para.paragraph_format.alignment
    newp.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
    newp.paragraph_format.keep_together = para.paragraph_format.keep_together
    newp.paragraph_format.keep_with_next = para.paragraph_format.keep_with_next
    newp.paragraph_format.left_indent = para.paragraph_format.left_indent
    newp.paragraph_format.line_spacing = para.paragraph_format.line_spacing
    newp.paragraph_format.line_spacing_rule = para.paragraph_format.line_spacing_rule
    newp.paragraph_format.page_break_before = para.paragraph_format.page_break_before
    newp.paragraph_format.right_indent = para.paragraph_format.right_indent
    newp.paragraph_format.space_after = para.paragraph_format.space_after
    newp.paragraph_format.space_before = para.paragraph_format.space_before
    for ts in para.paragraph_format.tab_stops:
        newp.paragraph_format.tab_stops.add_tab_stop(ts.position, ts.alignment, ts.leader)
    newp.paragraph_format.widow_control = para.paragraph_format.widow_control


def insert_paragraph_copy_before(doc, paraBefore, para):
    if paraBefore is None:
        newp = doc.add_paragraph()
    else:
        newp = paraBefore.insert_paragraph_before()
    copyPara(para, newp)
    for run in para.runs:
        add_run_copy(newp, run)
    return newp


def insert_paragraph_before(paraBefore, text, para):
    newp = paraBefore.insert_paragraph_before(text, para.style)
    copyPara(para, newp)
    return newp


def eqFont(f1, f2):
    if f1.name != f2.name:
        return False
    if f1.size != f2.size:
        return False
    return True


def eqStyle(s1, s2):
    if s1.name != s2.name:
        return False
    return True


def eqColor(r1, r2):
    p1 = hasattr(r1._element, "rPr")
    p2 = hasattr(r2._element, "rPr")
    if not p1 and not not p2:
        return True
    if p1 and not p2:
        return False
    if not p1 and p2:
        return False
    p1 = hasattr(r1._element.rPr, "color")
    p2 = hasattr(r2._element.rPr, "color")
    if not p1 and not p2:
        return True
    if p1 and not p2:
        return False
    if not p1 and p2:
        return False
    try:
        c1 = r1._element.rPr.color
        c2 = r2._element.rPr.color
        if c1 is None and c2 is None:
            return True
        if c1 is not None and c2 is None:
            return False
        if c1 is None and c2 is not None:
            return False
        return c1.val == c2.val
    except:
        logger.exception("eqcolor")


def split_run(para, runs, run, x):
    runX = runs.index(run) + 1
    t1 = run.text[0:x]
    t2 = run.text[x:]
    run.text = t1
    _ = add_run_copy(para, run, text=t2)
    # the insert does not work as expected, the new_run is always inserted into the same place,
    # irrespective of runX
    #    para._p.insert(runX, new_run._r)
    # therefore, we append all runs behind runX AFTER the newly appended run
    # i.e. we copy a b t1 c d t2 to a b t1 t2 c d, by appending c, d
    # this is all trial and error, and completely obscure...
    while runX < len(runs):
        para._p.append(runs[runX]._r)
        runX += 1
    if debug:
        print("splitRes:", " ".join(["<" + run.text + ">" for run in para.runs]))


"""
    This function combines the texts of successive runs with same 
    font,style,color into one run. Word splits for unknown reasons continuous 
    text like "Kommentar" into two runs "K" and "ommentar"!?
    Our parameters ${name} are split int "${", "name", "}". This makes 
    parsing too difficult, so we combine first.
    But then we may have several ${param}s within one run. We then split 
    the runs again so that each parameter is in its own run.
"""


def combineRuns(doc):
    paras = doc.paragraphs
    for para in paras:
        if debug:
            print("Para ", str(para), para.text, " align:", para.alignment,
                  "style:", para.style.name)
        runs = para.runs
        prevRun = None
        for run in runs:
            if debug:
                print("Run '", run.text, "' bold:", run.bold,
                      " font:", run.font.name, run.font.size,
                      " style:", run.style.name)
                # print("len ", len(run.text), " hex ", str2hex(run.text))
            if prevRun is not None and prevRun.bold == run.bold and \
                    prevRun.italic == run.italic and \
                    prevRun.underline == run.underline and \
                    eqColor(prevRun, run) and \
                    eqFont(prevRun.font, run.font) and \
                    eqStyle(prevRun.style, run.style) and \
                    run.text != "":
                prevRun.text += run.text
                delete_run(run)
            else:
                prevRun = run
    paras = doc.paragraphs
    for para in paras:
        if para.text.find("${") > 0:
            splitted = True
            while splitted:
                splitted = False
                runs = para.runs
                for run in runs:
                    mp = paramRE.search(run.text, 1)
                    if mp is None:
                        continue
                    sp = mp.span()
                    split_run(para, runs, run, sp[0])
                    splitted = True
                    break


def add_hyperlink_into_run(paragraph, run, i, url):
    runs = paragraph.runs
    if i is None:
        for i, runi in enumerate(runs):
            if run.text == runi.text:
                break
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
                          is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    hyperlink.append(run._r)
    # see above comment about insert
    # paragraph._p.insert(i+1, hyperlink)
    paragraph._p.append(hyperlink)
    i += 1
    while i < len(runs):
        paragraph._p.append(runs[i]._r)
        i += 1


def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()

    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
                              'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
                              'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
                              'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
                              'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
                              'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
                              'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
                              'w:pPrChange'
                              )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)


def move_run_before(i, para):
    runs = para.runs
    l = len(runs) - 1
    while i < l:
        para._p.append(runs[i]._r)
        i += 1


class DocxTreeHandler(markdown.treeprocessors.Treeprocessor):
    def __init__(self, md):
        super().__init__(md)
        self.curPara = None
        self.curRun = None
        self.ancestors = []
        self.states = []
        self.nodeHandler = {
            "h1": self.h1,
            "h2": self.h2,
            "h3": self.h3,
            "h4": self.h4,
            "h5": self.h5,
            "h6": self.h6,
            "p": self.p,
            "strong": self.strong,
            "em": self.em,
            "code": self.code,
            "blockquote": self.blockQuote,
            "stroke": self.stroke,
            "underline": self.underline,
            "ul": self.ul,
            "ol": self.ol,
            "li": self.li,
            "a": self.a,
            "img": self.img,
            "hr": self.hr}

    def run(self, root):
        self.paraBefore = self.docxHandler.para  # we insert everything before this paragraph
        self.para = copy.deepcopy(self.paraBefore)
        # self.para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # self.para.style = "Normal"
        self.curPara = None
        self.lvl = 4
        self.fontStyles = ""
        for child in root:  # skip <div> root
            self.walkOuter(child)
        root.clear()

    def setDeps(self, docxHandler):
        self.docxHandler = docxHandler

    @staticmethod
    def unescape(m):
        return chr(int(m.group(1)))

    def printLines(self, s):
        s = stxEtxRE.sub(self.unescape, s)  # "STX40ETX" -> chr(40), see markdown/postprocessors/UnescapePostprocessor
        r = self.curPara.add_run(s)  # style?
        r.bold = r.italic = r.font.strike = r.font.underline = False
        for fst in self.fontStyles:
            if fst == 'B':
                r.bold = True
            elif fst == 'I':
                r.italic = True
            elif fst == 'X':
                r.font.strike = True
            elif fst == 'U':
                r.font.underline = True
            elif fst == 'C':
                r.font.name = "Courier"
        self.curRun = r

    def walkOuter(self, node):
        global nlctr
        if debug:
            if node.text is not None:
                ltext = node.text.replace("\n", "<" + str(nlctr) + "nl>")
                # node.text = node.text.replace("\n", str(nlctr) + "\n")
                nlctr += 1
            else:
                ltext = "None"
            if node.tail is not None:
                ltail = node.tail.replace("\n", "<" + str(nlctr) + "nl>")
                # node.tail = node.tail.replace("\n", str(nlctr) + "\n")
                nlctr += 1
            else:
                ltail = "None"
            self.lvl += 4
            print(" " * self.lvl, "<<<<")
            print(" " * self.lvl, "node=", node.tag, ",text=", ltext,
                  "tail=", ltail)
        try:
            self.nodeHandler[node.tag](node)
            if node.tail is not None:
                self.printLines(node.tail)
        except Exception:
            msg = "Fehler während der Behandlung der Beschreibung des Events " + \
                  self.docxHandler.eventMsg
            logger.exception(msg)
            print(msg)
        if debug:
            print(" " * self.lvl, ">>>>")
            self.lvl -= 4

    def walkInner(self, node):
        if node.text is not None:
            self.printLines(node.text)
        for dnode in node:
            self.walkOuter(dnode)

    def h1(self, node):
        node.tail = None
        sav = self.para.style
        self.para.style = "Heading 1"
        self.curPara = insert_paragraph_before(self.paraBefore, None, self.para)
        self.walkInner(node)
        self.para.style = sav

    def h2(self, node):
        node.tail = None
        sav = self.para.style
        self.para.style = "Heading 2"
        self.curPara = insert_paragraph_before(self.paraBefore, None, self.para)
        self.walkInner(node)
        self.para.style = sav

    def h3(self, node):
        node.tail = None
        sav = self.para.style
        self.para.style = "Heading 3"
        self.curPara = insert_paragraph_before(self.paraBefore, None, self.para)
        self.walkInner(node)
        self.para.style = sav

    def h4(self, node):
        node.tail = None
        sav = self.para.style
        self.para.style = "Heading 4"
        self.curPara = insert_paragraph_before(self.paraBefore, None, self.para)
        self.walkInner(node)
        self.para.style = sav

    def h5(self, node):
        node.tail = None
        sav = self.para.style
        self.para.style = "Heading 5"
        self.curPara = insert_paragraph_before(self.paraBefore, None, self.para)
        self.walkInner(node)
        self.para.style = sav

    def h6(self, node):
        node.tail = None
        sav = self.para.style
        self.para.style = "Heading 6"
        self.curPara = insert_paragraph_before(self.paraBefore, None, self.para)
        self.walkInner(node)
        self.para.style = sav

    def p(self, node):
        node.tail = None
        sav = self.para.style
        self.para.style = "Normal"
        self.curPara = insert_paragraph_before(self.paraBefore, None, self.para)
        self.walkInner(node)
        self.para.style = sav

    def strong(self, node):
        sav = self.fontStyles
        self.fontStyles += "B"
        self.walkInner(node)
        self.fontStyles = sav

    def stroke(self, node):
        sav = self.fontStyles
        self.fontStyles += "X"
        self.walkInner(node)
        self.fontStyles = sav

    def underline(self, node):
        sav = self.fontStyles
        self.fontStyles += "U"
        self.walkInner(node)
        self.fontStyles = sav

    def em(self, node):
        sav = self.fontStyles
        self.fontStyles += "I"
        self.walkInner(node)
        self.fontStyles = sav

    def code(self, node):
        sav = self.fontStyles
        self.fontStyles += "C"
        self.walkInner(node)
        self.fontStyles = sav

    def ul(self, node):
        node.text = node.tail = None
        sav = self.para.style
        if self.para.style.name == "List Bullet":
            self.para.style = "List Bullet 2"
        else:
            self.para.style = "List Bullet"
        self.walkInner(node)
        self.para.style = sav

    def ol(self, node):
        node.text = node.tail = None
        sav = self.para.style
        if self.para.style.name == "List Number":
            self.para.style = "List Number 2"
        else:
            self.para.style = "List Number"
        self.walkInner(node)
        self.para.style = sav

    def li(self, node):
        node.tail = None
        self.curPara = insert_paragraph_before(self.paraBefore, None, self.para)
        self.walkInner(node)

    def a(self, node):
        url = node.attrib["href"]
        self.walkInner(node)
        add_hyperlink_into_run(self.curPara, self.curRun, None, url)
        self.curRun.font.color.rgb = RGBColor(238, 126, 13)

    def blockQuote(self, node):
        node.text = node.tail = None
        savA = self.para.alignment
        self.para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        self.walkInner(node)
        self.para.alignment = savA

    def hr(self, node):
        node.tail = None
        self.curPara = insert_paragraph_before(self.paraBefore, None, self.para)
        insertHR(self.curPara)
        self.walkInner(node)

    def img(self, node):
        self.walkInner(node)


class DocxExtension(markdown.Extension):
    def extendMarkdown(self, md):
        self.docxTreeHandler = DocxTreeHandler(md)
        md.treeprocessors.register(self.docxTreeHandler, "docxtreehandler", 5)
        md.inlinePatterns.register(
            markdown.inlinepatterns.SimpleTagInlineProcessor(
                strokeRE, 'stroke'), 'stroke', 40)
        md.inlinePatterns.register(
            markdown.inlinepatterns.SimpleTagInlineProcessor(
                ulRE, 'underline'), 'underline', 41)


class DocxHandler(expand.Expand):
    def __init__(self, gui):
        super().__init__()
        self.gui = gui
        self.terminselections = {}
        self.tourselections = {}
        self.touren = []
        self.termine = []
        self.url = None
        self.para = None
        self.run = None
        self.ausgabedatei = None
        self.runX = 0
        self.images = []
        self.selecter = selektion.Selektion()
        global debug
        try:
            _ = os.environ["DEBUG"]
            debug = True
        except:
            debug = False
        self.docxExtension = DocxExtension()
        self.md = markdown.Markdown(extensions=[self.docxExtension])
        self.docxExtension.docxTreeHandler.setDeps(self)

    def openDocx(self, pp):
        self.doc = docx.Document(self.gui.docxTemplateName)
        combineRuns(self.doc)
        self.parseParams()
        if pp:
            if debug:
                logger.debug("Styles: " + ", ".join([style.name for style in self.doc.styles]))
            self.setGuiParams()
        doc_styles = self.doc.styles
        wd2_style = doc_styles.add_style("WD2_STYLE", WD_STYLE_TYPE.CHARACTER)
        wd2_font = wd2_style.font
        wd2_font.name = "Wingdings 2"
        wd3_style = doc_styles.add_style("WD3_STYLE", WD_STYLE_TYPE.CHARACTER)
        wd3_font = wd3_style.font
        wd3_font.name = "Wingdings 3"

    @staticmethod
    def nothingFound():
        logger.info("Nichts gefunden")
        print("Nichts gefunden")

    def parseParams(self):
        texts = []
        for para in self.doc.paragraphs:
            delete_paragraph(para)
            if para.text.find("-----") >= 0:
                break
            if para.text == "" or para.text.startswith("Kommentar"):
                continue
            if para.style.name.startswith("List"):
                continue
            if para.paragraph_format.left_indent is not None:
                continue
            texts.append(para.text)
        lines = "\n".join(texts).split('\n')
        # defaults:
        self.linkType = "Frontend"
        self.includeSub = True
        lx = 0
        selections = {}
        while lx < len(lines):
            line = lines[lx]
            words = line.split()
            if len(words) == 0:
                lx += 1
                continue
            word0 = words[0].lower().replace(":", "")
            if len(words) > 1:
                if word0 == "linktyp":
                    self.linkType = words[1].lower().capitalize()
                    lx += 1
                elif word0 == "ausgabedatei":
                    self.ausgabedatei = words[1]
                    lx += 1
                else:
                    raise ValueError(
                        "Unbekannter Parameter " + word0 +
                        ", erwarte linktyp oder ausgabedatei")
            elif word0 not in ["selektion", "terminselektion", "tourselektion"]:
                raise ValueError(
                    "Unbekannter Parameter " + word0 +
                    ", erwarte selektion, terminselektion oder tourselektion")
            else:
                lx = self.parseSel(word0, lines, lx + 1, selections)

        selection = selections.get("selektion")
        self.gliederung = selection.get("gliederungen")
        self.includeSub = selection.get("mituntergliederungen") == "ja"
        self.start = selection.get("beginn")
        self.end = selection.get("ende")

        sels = selections.get("terminselektion")
        if sels is not None:
            for sel in sels.values():
                self.terminselections[sel.get("name")] = sel
                for key in sel.keys():
                    if key != "name" and not isinstance(sel[key], list):
                        sel[key] = [sel[key]]

        sels = selections.get("tourselektion")
        if sels is not None:
            for sel in sels.values():
                self.tourselections[sel.get("name")] = sel
                for key in sel.keys():
                    if key != "name" and not isinstance(sel[key], list):
                        sel[key] = [sel[key]]

    def setGuiParams(self):
        self.gui.setLinkType(self.linkType)
        if self.gliederung is not None and self.gliederung != "":
            self.gui.setGliederung(self.gliederung)
        self.gui.setIncludeSub(self.includeSub)
        if self.start is not None and self.start != "":
            self.gui.setStart(self.start)
        if self.end is not None and self.end != "":
            self.gui.setEnd(self.end)
        self.setEventType()
        self.setRadTyp()

    def setEventType(self):
        typ = ""
        if len(self.terminselections) != 0 and len(self.tourselections) != 0:
            typ = "Alles"
        elif len(self.terminselections) != 0:
            typ = "Termin"
        elif len(self.tourselections) != 0:
            typ = "Radtour"
        if typ != "":
            self.gui.setEventType(typ)

    def setRadTyp(self):
        rts = set()
        for sel in self.tourselections.values():
            l = sel.get("radtyp")
            if l is None or len(l) == 0:
                l = [self.gui.getRadTyp()]
            for elem in l:
                rts.add(elem)
        if "Alles" in rts:
            typ = "Alles"
        elif len(rts) == 1:
            typ = rts.pop()
        else:
            typ = "Alles"
        self.gui.setRadTyp(typ)

    def getIncludeSub(self):
        return self.includeSub

    def getEventType(self):
        if len(self.terminselections) != 0 and len(self.tourselections) != 0:
            return "Alles"
        if len(self.terminselections) != 0:
            return "Termin"
        if len(self.tourselections) != 0:
            return "Radtour"
        return self.gui.getEventType()

    def getRadTyp(self):
        rts = set()
        for sel in self.tourselections.values():
            l = sel.get("radtyp")
            if l is None or len(l) == 0:
                l = [self.gui.getRadTyp()]
            for elem in l:
                rts.add(elem)
        if "Alles" in rts:
            return "Alles"
        if len(rts) == 1:
            return rts[0]
        return "Alles"

    def getUnitKeys(self):
        return self.gliederung

    def getStart(self):
        return self.start

    def getEnd(self):
        return self.end

    @staticmethod
    def parseSel(word, lines, lx, selections):
        selections[word] = sel = sel2 = {}
        while lx < len(lines):
            line = lines[lx]
            if line.strip() == "":
                lx += 1
                continue
            if not line[0].isspace():
                return lx
            words = line.split()
            word0 = words[0].lower().replace(":", "")
            if word0 == "name":
                word1 = words[1].lower()
                sel[word1] = sel2 = {}
                sel2["name"] = word1
            else:
                lst = ",".join(words[1:]).split(",")
                sel2[word0] = lst[0] if len(lst) == 1 else lst
            lx += 1
        return lx

    def handleTour(self, tour):
        self.touren.append(tour)

    def handleTermin(self, termin):
        self.termine.append(termin)

    def handleEnd(self):
        print("Template", self.gui.docxTemplateName, "wird abgearbeitet")
        if self.doc is None:
            self.openDocx(False)
        self.linkType = self.gui.getLinkType()
        self.gliederung = self.gui.getGliederung()
        self.includeSub = self.gui.getIncludeSub()
        self.start = self.gui.getStart()
        self.end = self.gui.getEnd()
        self.images = []
        paragraphs = self.doc.paragraphs
        paraCnt = len(paragraphs)
        paraNo = 0
        while paraNo < paraCnt:
            para = paragraphs[paraNo]
            if para.text.startswith("/template"):
                p1 = paraNo
                while True:
                    if para.text.find("/endtemplate") >= 0:
                        break
                    delete_paragraph(para)
                    paraNo += 1
                    para = paragraphs[paraNo]
                delete_paragraph(para)
                p2 = paraNo
                paraNo += 1
                self.paraBefore = \
                    None if paraNo == paraCnt else paragraphs[paraNo]
                tempParas = paragraphs[p1:p2 + 1]
                self.evalTemplate(tempParas)
            else:
                self.evalPara(para)
                paraNo += 1

        ausgabedatei = self.ausgabedatei
        if ausgabedatei is None or ausgabedatei == "":
            ausgabedatei = os.path.join(
                os.path.dirname(self.gui.docxTemplateName),
                "ADFC_" +
                self.gliederung +
                ("_I_" if self.includeSub else "_") +
                self.start + "-" +
                self.end + "_" +
                self.linkType[0] +
                ".docx")
        try:
            self.doc.save(ausgabedatei)
            print("Ausgabedatei", ausgabedatei, "wurde erzeugt")
            try:
                opath = os.path.abspath(ausgabedatei)
                os.startfile(opath)
            except Exception:
                logger.exception("opening " + ausgabedatei)
        except Exception as e:
            print("Ausgabedatei", ausgabedatei, "konnte nicht geschrieben werden")
            raise e
        finally:
            self.doc = None
            self.touren = []
            self.termine = []

    def evalPara(self, para):
        if debug:
            print("para", para.text)
        for run in para.runs:
            self.evalRun(run, None)

    def evalTemplate(self, paras):
        if debug:
            print("template:")
        para0 = paras[0]
        para0Lines = para0.runs[0].text.split('\n')
        if not para0Lines[0].startswith("/template"):
            raise ValueError("/template muss am Anfang der ersten Zeile eines Paragraphen stehen")
        paraN = paras[-1]
        paraNLines = paraN.runs[-1].text.split('\n')
        if not paraNLines[-1].startswith("/endtemplate"):
            raise ValueError("/endtemplate muss am Anfang der letzten Zeile eines Paragraphen stehen")
        words = para0Lines[0].split()
        typ = words[1]
        if typ != "/tour" and typ != "/termin":
            raise ValueError("Zweites Wort nach /template muß /tour oder /termin sein")
        typ = typ[1:]
        sel = words[2]
        if not sel.startswith("/selektion="):
            raise ValueError("Drittes Wort nach /template muß mit /selektion= beginnen")
        sel = sel[11:].lower()
        sels = self.tourselections if typ == "tour" else self.terminselections
        if sel not in sels:
            raise ValueError("Selektion " + sel + " nicht in " + typ + "selektion")
        sel = sels[sel]
        events = self.touren if typ == "tour" else self.termine
        self.evalEvents(sel, events, paras)

    def evalEvents(self, sel, events, paras):
        selectedEvents = []
        for event in events:
            if self.selecter.selected(event, sel):
                selectedEvents.append(event)
        if len(selectedEvents) == 0:
            return
        for event in selectedEvents:
            self.eventMsg = event.getTitel() + " vom " + event.getDatum()[0]
            for para in paras:
                newp = insert_paragraph_copy_before(
                    self.doc, self.paraBefore, para)
                self.para = newp
                for self.runX, self.run in enumerate(newp.runs):
                    run = self.run
                    if run.text.lower().startswith("/kommentar"):
                        continue
                    rtext = run.text.strip()
                    self.evalRun(run, event)
                    if rtext == "${titel}" and self.url is not None:
                        add_hyperlink_into_run(newp, run, self.runX, self.url)
                        # newp.add_run().add_picture(io.BytesIO(base64.decodebytes(event.getImagePreview().encode())))
                        if self.gui.getIncludeImg():
                            try:
                                image = event.getImageStream(event.getImageUrl(), event.getEventItemId())
                                self.images.append(image) # see adfc_gui.py
                                newp.add_run().add_picture(image, width=4000000.0)
                            except Exception as e:
                                logger.exception("cannot get image")
                if newp.text == "":
                    delete_paragraph(newp)

    def evalRun(self, run, event):
        if debug:
            print("run", run.text)
        linesOut = []
        linesIn = run.text.split('\n')
        for line in linesIn:
            if not line.startswith("/template") and \
                    not line.startswith("/endtemplate"):
                exp = self.expand(line, event)
                if exp is not None:
                    linesOut.append(exp)
        newtext = '\n'.join(linesOut)
        if run.text != newtext:
            if run.text.lower() == "${schwierigkeitm}":
                self.para.add_run(text=newtext, style="WD2_STYLE")
                move_run_before(self.runX, self.para)
                delete_run(run)
            elif run.text.lower() == "${schwierigkeith}":
                self.para.add_run(text=newtext, style="WD3_STYLE")
                move_run_before(self.runX, self.para)
                delete_run(run)
            else:
                run.text = newtext  # assignment to run.text lets images disappear!?!?

    def expBeschreibung(self, event, _):
        if len(self.para.runs) != 1:
            raise ValueError("${beschreibung} muß in einem Paragraphen einzeln stehen")
        desc = event.getBeschreibung(True)
        # desc = codecs.decode(desc, encoding = "unicode_escape")
        self.md.convert(desc)
        self.md.reset()
        return None

    def expPersonen(self, bezeichnung, event):
        tl = event.getPersonen()
        if len(tl) == 0:
            return ""

        # print("TL0:", self.runX, "<<" + self.para.runs[self.runX].text + ">>", " ".join(["<" + run.text + ">" for run in self.para.runs]))
        run = self.para.add_run(text=bezeichnung + ": ", style=self.run.style)
        run.bold = True
        move_run_before(self.runX, self.para)
        # print("TL1:", " ".join(["<" + run.text + ">" for run in self.para.runs]))
        self.runX += 1

        self.para.add_run(text=", ".join(tl), style=self.run.style)
        move_run_before(self.runX, self.para)
        # print("TL2:", " ".join(["<" + run.text + ">" for run in self.para.runs]))
        return ""

    def expAbfahrten(self, tour, _):
        afs = tour.getAbfahrten()
        if len(afs) == 0:
            return ""
        afl = []
        for af in afs:
            if af[1] == "":
                afl.append(af[2])
            else:
                afl.append(af[0] + " " + af[1] + " " + af[2])
        # print("AB0:", self.runX, "<<" + self.para.runs[self.runX].text + ">>", " ".join(["<" + run.text + ">" for run in self.para.runs]))

        run = self.para.add_run(text="Ort" + ("" if len(afs) == 1 else "e") + ": ", style=self.run.style)
        run.bold = True
        move_run_before(self.runX, self.para)
        # print("AB1:", " ".join(["<" + run.text + ">" for run in self.para.runs]))
        self.runX += 1

        self.para.add_run(text=", ".join(afl), style=self.run.style)
        move_run_before(self.runX, self.para)
        # print("AB2:", " ".join(["<" + run.text + ">" for run in self.para.runs]))

        return ""

    def expZusatzInfo(self, tour, _):
        zi = tour.getZusatzInfo()
        if len(zi) == 0:
            return None
        for z in zi:
            # print("ZU0:", self.runX, "<<" + self.para.runs[self.runX].text + ">>",
            # " ".join(["<" + run.text + ">" for run in self.para.runs]))
            x = z.find(':') + 1
            run = self.para.add_run(text=z[0:x], style=self.run.style)
            run.bold = True
            move_run_before(self.runX, self.para)
            # print("ZU1:", " ".join(["<" + run.text + ">" for run in self.para.runs]))
            self.runX += 1

            self.para.add_run(text=z[x:] + "\n", style=self.run.style)
            move_run_before(self.runX, self.para)
            # print("ZU2:", " ".join(["<" + run.text + ">" for run in self.para.runs]))
            self.runX += 1
        return ""
