import re

import docx
from lxml import etree

paramRE = re.compile(r"\${(\w*?)}")

def delete_paragraph(paragraph):
    # https://github.com/python-openxml/python-docx/issues/33
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def delete_run(run):
    r = run._element
    r.getparent().remove(r)
    r._parent = r._element = None

def add_run_copy(paragraph, run, text=None):
    r = paragraph.add_run(text=run.text if text is None else text, style=run.style)
    r.bold = run.bold
    r.italic = run.italic
    r.underline = run.underline
    r.font.all_caps = run.font.all_caps
    r.font.bold = run.font.bold
    r.font.color.rgb = run.font.color.rgb
    r.font.color.theme_color = run.font.color.theme_color
    #r.font.color.type = run.font.color.type
    r.font.complex_script = run.font.complex_script
    r.font.cs_bold = run.font.cs_bold
    r.font.cs_italic = run.font.cs_italic
    r.font.double_strike = run.font.double_strike
    r.font.emboss = run.font.emboss
    r.font.hidden = run.font.hidden
    r.font.highlight_color = run.font.highlight_color
    r.font.imprint = run.font.imprint
    r.font.italic = run.font.italic
    r.font.math = run.font.math
    r.font.name = run.font.name
    r.font.no_proof = run.font.no_proof
    r.font.outline = run.font.outline
    r.font.rtl = run.font.rtl
    r.font.shadow = run.font.shadow
    r.font.size = run.font.size
    r.font.small_caps = run.font.small_caps
    r.font.snap_to_grid = run.font.snap_to_grid
    r.font.spec_vanish = run.font.spec_vanish
    r.font.strike = run.font.strike
    r.font.subscript = run.font.subscript
    r.font.superscript = run.font.superscript
    r.font.underline = run.font.underline
    r.font.web_hidden = run.font.web_hidden
    return r


def insert_paragraph_copy_before(paraBefore, para):
    newp = paraBefore.insert_paragraph_before()
    newp.alignment = para.alignment
    newp.style = para.style
    newp.paragraph_format.alignment = para.paragraph_format.alignment
    newp.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
    newp.paragraph_format.keep_together = para.paragraph_format.keep_together
    newp.paragraph_format.keep_with_next = para.paragraph_format.keep_with_next
    newp.paragraph_format.left_indent = para.paragraph_format.left_indent
    newp.paragraph_format.line_spacing = para.paragraph_format.line_spacing
    newp.paragraph_format.line_spacing_rule = para.paragraph_format.line_spacing_rule
    newp.paragraph_format.page_break_before = para.paragraph_format.page_break_before
    newp.paragraph_format.right_indent = para.paragraph_format.right_indent
    newp.paragraph_format.space_after = para.paragraph_format.space_after
    newp.paragraph_format.space_before = para.paragraph_format.space_before
    for ts in para.paragraph_format.tab_stops:
        newp.paragraph_format.add_tab_stop(ts-position, ts.alignment, ts.leader)
    newp.paragraph_format.widow_control = para.paragraph_format.widow_control
    for run in para.runs:
        add_run_copy(newp, run)

def eqFont(f1, f2):
    if f1.name != f2.name:
        return False
    if f1.size != f2.size:
        return False
    return True

def eqStyle(s1, s2):
    if s1.name != s2.name:
        return False
    return True

def eqColor(r1, r2):
    p1 = hasattr(r1._element, "rPr")
    p2 = hasattr(r2._element, "rPr")
    if not p1 and not not p2:
        return True
    if p1 and not p2:
        return False
    if not p1 and p2:
        return False
    p1 = hasattr(r1._element.rPr, "color")
    p2 = hasattr(r2._element.rPr, "color")
    if not p1 and not p2:
        return True
    if p1 and not p2:
        return False
    if not p1 and p2:
        return False
    try:
        c1 = r1._element.rPr.color
        c2 = r2._element.rPr.color
    except:
        print("!!")
    if c1 == None and c2 == None:
        return True
    if c1 != None and c2 == None:
        return False
    if c1 == None and c2 != None:
        return False
    return c1.val == c2.val

def split_run(para, runs, run, x):
    runX = runs.index(run) + 1
    t1 = run.text[0:x]
    t2 = run.text[x:]
    run.text = t1
    new_run = add_run_copy(para, run, text=t2)
    # the insert does not work as expected, the new_run is always inserted into the same place,
    # irrespective of runX
    #    para._p.insert(runX, new_run._r)
    # therefore, we append all runs behind runX AFTER the newly appended run
    # i.e. we copy a b t1 c d t2 to a b t1 t2 c d, by appending c, d
    # this is all trial and error, and completely obscure...
    while runX < len(runs):
        para._p.append(runs[runX]._r)
        runX += 1
    print("splitRes:", " ".join(["<" + run.text + ">" for run in para.runs]))


def combineRuns(doc):
    paras = doc.paragraphs
    for para in paras:
        print("Para ", str(para), para.text, " align:", para.alignment, "style:", para.style.name)
        runs = para.runs
        prevRun = None
        for run in runs:
            print("Run '", run.text, "' bold:", run.bold, " font:", run.font.name, run.font.size, " style:", run.style.name)
            if prevRun != None and prevRun.bold == run.bold and prevRun.italic == run.italic and \
                prevRun.underline == run.underline and \
                eqColor(prevRun, run) and \
                eqFont(prevRun.font, run.font) and \
                eqStyle(prevRun.style, run.style):
                prevRun.text += run.text
                delete_run(run)
            else:
                prevRun = run
    paras = doc.paragraphs
    for para in paras:
        if para.text.find("${") > 0:
            splitted = True
            while splitted:
                splitted = False
                runs = para.runs
                for run in runs:
                    mp = paramRE.search(run.text, 1)
                    if mp == None:
                        continue
                    sp = mp.span()
                    split_run(para, runs, run, sp[0])
                    splitted = True
                    break

def repeat(doc, n):
    paras = doc.paragraphs
    lastpara = paras[-1]
    for para in paras[-n:-1]:
        insert_paragraph_copy_before(lastpara, para)

def add_hyperlink(paragraph, url, text):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

class DocxTemplate(object):
    def __init__(self, docx):
        self.docx = docx
    def __getattr__(self, name):
        return getattr(self.docx, name)

    def xml_to_string(self, xml, encoding='unicode'):
        # Be careful : pretty_print MUST be set to False, otherwise patch_xml()
        # won't work properly
        return etree.tostring(xml, encoding='unicode', pretty_print=True)

    def get_docx(self):
        return self.docx

    def get_xml(self):
        return self.xml_to_string(self.docx._element.body)

def docread(sav):
    if sav:
        doc = docx.Document(r"c:/temp/sav.docx")
    else:
        #doc = docx.Document(r"c:/temp/test.docx")
        doc = docx.Document("template.docx")
    #dt1 = DocxTemplate(doc)
    #xml1 = dt1.get_xml()
    combineRuns(doc)
    #xml2 = dt1.get_xml()
    if not sav:
        #repeat(doc, 50)
        doc.save("c:/temp/sav.docx")

docread(False)



"""
document = docx.Document()
p = document.add_paragraph()
add_hyperlink(p, 'http://www.baidu.com', 'baidu')
document.save('demo_hyperlink.docx'
"""
